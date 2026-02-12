"""
Text Extractor Module
Extracts text content from PDF and DOCX files
"""

import pdfplumber
from docx import Document
from typing import Optional
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class TextExtractor:
    """
    A class to extract text from various document formats.
    
    Supports:
    - PDF files (using pdfplumber)
    - DOCX files (using python-docx)
    """
    
    @staticmethod
    def extract_from_pdf(file_path: str) -> Optional[str]:
        """
        Extract text content from a PDF file.
        
        Args:
            file_path (str): Path to the PDF file or file object
            
        Returns:
            Optional[str]: Extracted text or None if extraction fails
        """
        try:
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            if not text.strip():
                logger.warning("PDF appears to be empty or scanned (no extractable text)")
                return None
                
            logger.info(f"Successfully extracted {len(text)} characters from PDF")
            return text
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            return None
    
    @staticmethod
    def extract_from_docx(file_path: str) -> Optional[str]:
        """
        Extract text content from a DOCX file.
        
        Args:
            file_path (str): Path to the DOCX file or file object
            
        Returns:
            Optional[str]: Extracted text or None if extraction fails
        """
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\n"
            
            if not text.strip():
                logger.warning("DOCX appears to be empty")
                return None
                
            logger.info(f"Successfully extracted {len(text)} characters from DOCX")
            return text
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            return None
    
    @staticmethod
    def extract_text(file_path: str, file_type: str) -> Optional[str]:
        """
        Extract text from a file based on its type.
        
        Args:
            file_path (str): Path to the file
            file_type (str): Type of file ('pdf' or 'docx')
            
        Returns:
            Optional[str]: Extracted text or None if extraction fails
        """
        file_type = file_type.lower()
        
        if file_type == 'pdf':
            return TextExtractor.extract_from_pdf(file_path)
        elif file_type in ['docx', 'doc']:
            return TextExtractor.extract_from_docx(file_path)
        else:
            logger.error(f"Unsupported file type: {file_type}")
            return None


def extract_text(file_path: str, file_type: str) -> Optional[str]:
    """
    Convenience function to extract text from a file.
    
    Args:
        file_path (str): Path to the file
        file_type (str): Type of file ('pdf' or 'docx')
        
    Returns:
        Optional[str]: Extracted text or None if extraction fails
    """
    return TextExtractor.extract_text(file_path, file_type)
